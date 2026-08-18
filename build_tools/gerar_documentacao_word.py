"""Gera a documentação técnica consolidada do FAS Jornada em Word."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT = DOCS_DIR / "FAS_Jornada_Documentacao_Tecnica_v8.1.docx"
VERSION = "8.1"

CHAPTERS = (
    ("Visão geral", ROOT / "README.md"),
    ("Arquitetura", DOCS_DIR / "ARQUITETURA.md"),
    ("Referência dos módulos", DOCS_DIR / "REFERENCIA_MODULOS.md"),
    ("Fluxos de processamento", DOCS_DIR / "FLUXOS.md"),
    ("Integração RHiD", DOCS_DIR / "INTEGRACAO_RHID.md"),
    ("Integração Power BI", DOCS_DIR / "INTEGRACAO_POWER_BI.md"),
    ("Segurança e privacidade", DOCS_DIR / "SEGURANCA_E_PRIVACIDADE.md"),
    ("Desenvolvimento e testes", DOCS_DIR / "DESENVOLVIMENTO.md"),
    ("Instalador e empacotamento", DOCS_DIR / "INSTALADOR.md"),
)

BLUE = "2A495B"
LIGHT_BLUE = "DDEAF1"
GRAY = "5E6E77"


def _field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, text, separate, end))


def _set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _add_inline(paragraph, text: str) -> None:
    pattern = re.compile(
        r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`)"
    )
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        else:
            label, target = re.match(r"\[([^]]+)\]\(([^)]+)\)", token).groups()
            paragraph.add_run(f"{label} ({target})")
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _add_table(document: Document, lines: list[str]) -> None:
    rows = [
        [column.strip() for column in line.strip().strip("|").split("|")]
        for line in lines
    ]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", item) for item in rows[1]):
        rows.pop(1)
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        _set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(values[: len(cells)]):
            cells[index].text = value.replace("`", "")
    document.add_paragraph()


def _add_code(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.right_indent = Cm(0.6)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F5F7")
    properties.append(shading)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def _add_markdown(document: Document, path: Path, *, skip_title: bool = True) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading_skipped = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if stripped.startswith("```"):
            code: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            _add_code(document, code)
        elif stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines):
                current = lines[index].strip()
                if not (current.startswith("|") and current.endswith("|")):
                    break
                table_lines.append(current)
                index += 1
            _add_table(document, table_lines)
            continue
        elif stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if skip_title and level == 1 and not first_heading_skipped:
                first_heading_skipped = True
            else:
                document.add_heading(title, level=min(level + 1, 4))
        elif re.match(r"^[-*] ", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _add_inline(paragraph, re.sub(r"^\d+\. ", "", stripped))
        elif stripped and set(stripped) == {"-"}:
            pass
        elif stripped:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            _add_inline(paragraph, stripped)
        index += 1


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("263B47")
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Title", 28), ("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.font.bold = True


def _configure_sections(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("FAS Jornada • Documentação técnica • ")
        _field(footer, "PAGE")


def generate() -> Path:
    document = Document()
    _configure_styles(document)
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)

    logo = ROOT / "app" / "assets" / "logo.png"
    if logo.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo), width=Inches(2.2))

    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FAS Jornada")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documentação técnica completa")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(30)
    metadata.add_run(f"Versão {VERSION}\n")
    metadata.add_run(f"Gerado em {date.today().strftime('%d/%m/%Y')}\n")
    metadata.add_run("Relatório e Análise de Jornada")

    document.add_page_break()
    document.add_heading("Sumário", level=1)
    toc = document.add_paragraph()
    _field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph(
        "No Microsoft Word, clique com o botão direito no sumário e escolha "
        "“Atualizar Campo” para recalcular números de página."
    )
    note.runs[0].italic = True
    note.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

    for number, (title, path) in enumerate(CHAPTERS, start=1):
        document.add_page_break()
        document.add_heading(f"{number}. {title}", level=1)
        _add_markdown(document, path)

    document.add_section(WD_SECTION.CONTINUOUS)
    _configure_sections(document)
    document.core_properties.title = "FAS Jornada — Documentação técnica completa"
    document.core_properties.subject = "Arquitetura, módulos, fluxos, RHiD, segurança, testes e distribuição"
    document.core_properties.author = "FAS"
    document.core_properties.keywords = "FAS Jornada, RHiD, relatório, jornada, documentação"
    document.core_properties.comments = "Gerado a partir da documentação versionada do projeto."
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(generate())
